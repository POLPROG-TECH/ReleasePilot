"""Tests for executive brief composition and rendering.

Covers:
- ExecutiveBrief model and composition logic
- Business-language title transformation
- Executive summary generation
- Impact area grouping
- Risk extraction and next-steps generation
- Executive Markdown, PDF, and DOCX renderers
- CLI integration for executive audience
"""

from __future__ import annotations

import json
from datetime import date

import pytest

from releasepilot.audience.executive import (
    ExecutiveBrief,
    compose_executive_brief,
)
from releasepilot.domain.enums import ChangeCategory, Importance
from releasepilot.domain.models import (
    ChangeGroup,
    ChangeItem,
    ReleaseNotes,
    ReleaseRange,
    SourceReference,
)
from releasepilot.rendering.executive_md import ExecutiveMarkdownRenderer


def _pdf_available() -> bool:
    try:
        import reportlab  # noqa: F401

        return True
    except ImportError:
        return False


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


# ── Fixtures ─────────────────────────────────────────────────────────────────


def _make_item(
    title: str,
    category: ChangeCategory = ChangeCategory.FEATURE,
    *,
    breaking: bool = False,
    importance: Importance = Importance.NORMAL,
    description: str = "",
    scope: str = "",
) -> ChangeItem:
    return ChangeItem(
        id=f"id-{title[:10]}",
        title=title,
        description=description,
        category=category,
        scope=scope,
        importance=importance,
        is_breaking=breaking,
        source=SourceReference(commit_hash="abc12345"),
        authors=("alice",),
    )


@pytest.fixture()
def sample_notes() -> ReleaseNotes:
    """ReleaseNotes with a realistic mix of changes."""
    items_by_cat = {
        ChangeCategory.FEATURE: [
            _make_item("feat: add OAuth2 authentication", description="Full OAuth2 flow with PKCE"),
            _make_item("feat: add dark mode support"),
            _make_item("feat: implement search API"),
        ],
        ChangeCategory.BUGFIX: [
            _make_item("fix: pagination off-by-one error", ChangeCategory.BUGFIX),
            _make_item("fix: session token refresh", ChangeCategory.BUGFIX),
        ],
        ChangeCategory.SECURITY: [
            _make_item(
                "fix: patch XSS vulnerability", ChangeCategory.SECURITY, importance=Importance.HIGH
            ),
        ],
        ChangeCategory.PERFORMANCE: [
            _make_item("perf: optimize dashboard queries", ChangeCategory.PERFORMANCE),
        ],
        ChangeCategory.IMPROVEMENT: [
            _make_item("improve error messages", ChangeCategory.IMPROVEMENT),
        ],
        ChangeCategory.BREAKING: [
            _make_item(
                "feat!: remove /api/v1 endpoints",
                ChangeCategory.BREAKING,
                breaking=True,
                description="All v1 API endpoints have been removed",
            ),
        ],
        ChangeCategory.DEPRECATION: [
            _make_item("deprecate legacy webhook format", ChangeCategory.DEPRECATION),
        ],
    }

    groups = tuple(
        ChangeGroup(category=cat, items=tuple(items)) for cat, items in items_by_cat.items()
    )

    all_items = [i for items in items_by_cat.values() for i in items]
    highlights = tuple(i for i in all_items if i.is_breaking or i.importance == Importance.HIGH)
    breaking = tuple(i for i in all_items if i.is_breaking)

    return ReleaseNotes(
        release_range=ReleaseRange(
            from_ref="v2.0.0",
            to_ref="v3.0.0",
            version="3.0.0",
            title="Release 3.0.0",
            release_date=date(2026, 3, 16),
        ),
        groups=groups,
        highlights=highlights,
        breaking_changes=breaking,
        total_changes=len(all_items),
    )


@pytest.fixture()
def sample_brief(sample_notes: ReleaseNotes) -> ExecutiveBrief:
    return compose_executive_brief(sample_notes)


# ── ExecutiveBrief composition tests ─────────────────────────────────────────


# ── Executive Markdown renderer tests ────────────────────────────────────────


# ── Executive PDF renderer tests ─────────────────────────────────────────────


# ── Executive DOCX renderer tests ────────────────────────────────────────────


# ── CLI integration tests ────────────────────────────────────────────────────


# ── Audience view tests ──────────────────────────────────────────────────────


class TestExecutiveMarkdownRenderer:
    """Scenarios for rendering an executive brief to Markdown."""

    """GIVEN an executive brief"""

    def test_contains_executive_summary_heading(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN it contains the Executive Summary heading"""
        assert "## Executive Summary" in output

    """GIVEN an executive brief"""

    def test_contains_key_achievements(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN it contains the Key Achievements heading"""
        assert "## Key Achievements" in output

    """GIVEN an executive brief"""

    def test_contains_risks_section(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN it contains the Risks section"""
        assert "## Risks & Attention Items" in output

    """GIVEN an executive brief"""

    def test_contains_next_steps(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN it contains the Next Steps section"""
        assert "## Recommended Next Steps" in output

    """GIVEN an executive brief"""

    def test_contains_metrics_table(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN it contains the metrics table"""
        assert "## Release Metrics" in output
        assert "Total Changes" in output

    """GIVEN an executive brief"""

    def test_contains_report_title(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN it contains the report title with version"""
        assert "Release Brief" in output
        assert "3.0.0" in output

    """GIVEN an executive brief"""

    def test_no_conventional_commit_prefixes(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown"""

        """THEN no conventional commit prefixes appear"""
        # Should not contain raw commit prefixes
        assert "feat:" not in output
        assert "fix:" not in output
        assert "perf:" not in output


class TestExecutiveJsonRenderer:
    """Scenarios for rendering an executive brief to JSON."""

    """GIVEN an executive brief"""

    def test_valid_json(self, sample_brief: ExecutiveBrief):
        output = ExecutiveMarkdownRenderer().render_json(sample_brief)

        """WHEN rendered to JSON"""
        data = json.loads(output)

        """THEN the output is valid JSON with correct type"""
        assert data["type"] == "executive_brief"

    """GIVEN an executive brief"""

    def test_json_has_all_sections(self, sample_brief: ExecutiveBrief):
        """WHEN rendered to JSON"""
        data = json.loads(ExecutiveMarkdownRenderer().render_json(sample_brief))

        """THEN all expected sections are present"""
        assert "executive_summary" in data
        assert "key_achievements" in data
        assert "impact_areas" in data
        assert "risks" in data
        assert "next_steps" in data
        assert "metrics" in data

    """GIVEN an executive brief"""

    def test_json_metrics_match(self, sample_brief: ExecutiveBrief):
        """WHEN rendered to JSON"""
        data = json.loads(ExecutiveMarkdownRenderer().render_json(sample_brief))

        """THEN metrics match the original brief"""
        assert data["metrics"]["total_changes"] == 10
        assert data["metrics"]["features"] == 3


@pytest.mark.skipif(not _pdf_available(), reason="reportlab not installed")
class TestExecutivePdfRenderer:
    """Scenarios for rendering an executive brief to PDF."""

    """GIVEN an executive brief"""

    def test_produces_pdf_bytes(self, sample_brief: ExecutiveBrief):
        from releasepilot.rendering.executive_pdf import ExecutivePdfRenderer

        """WHEN rendered to PDF"""
        data = ExecutivePdfRenderer().render_bytes(sample_brief)

        """THEN valid PDF bytes are produced"""
        assert data[:5] == b"%PDF-"
        assert len(data) > 500

    """GIVEN an executive brief"""

    def test_pdf_has_meaningful_size(self, sample_brief: ExecutiveBrief):
        from releasepilot.rendering.executive_pdf import ExecutivePdfRenderer

        """WHEN rendered to PDF"""
        data = ExecutivePdfRenderer().render_bytes(sample_brief)

        """THEN the output has meaningful size"""
        assert len(data) > 2000


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
class TestExecutiveDocxRenderer:
    """Scenarios for rendering an executive brief to DOCX."""

    """GIVEN an executive brief"""

    def test_produces_docx_bytes(self, sample_brief: ExecutiveBrief):
        from releasepilot.rendering.executive_docx import ExecutiveDocxRenderer

        """WHEN rendered to DOCX"""
        data = ExecutiveDocxRenderer().render_bytes(sample_brief)

        """THEN valid DOCX bytes are produced"""
        assert data[:2] == b"PK"  # DOCX is a ZIP archive
        assert len(data) > 1000

    """GIVEN an executive brief"""

    def test_docx_reopens_as_document(self, sample_brief: ExecutiveBrief):
        from io import BytesIO

        from docx import Document

        from releasepilot.rendering.executive_docx import ExecutiveDocxRenderer

        """WHEN rendered to DOCX and reopened"""
        data = ExecutiveDocxRenderer().render_bytes(sample_brief)
        doc = Document(BytesIO(data))

        """THEN it contains multiple paragraphs"""
        # Should have multiple paragraphs
        assert len(doc.paragraphs) > 5

    """GIVEN an executive brief"""

    def test_docx_contains_executive_summary(self, sample_brief: ExecutiveBrief):
        from io import BytesIO

        from docx import Document

        from releasepilot.rendering.executive_docx import ExecutiveDocxRenderer

        """WHEN rendered to DOCX"""
        data = ExecutiveDocxRenderer().render_bytes(sample_brief)
        doc = Document(BytesIO(data))

        """THEN the document contains Executive Summary text"""
        # Collect text from paragraphs AND table cells
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts)
        assert "Executive Summary" in text

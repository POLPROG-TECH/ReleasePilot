"""Tests for executive brief composition and rendering.

Covers:
- ExecutiveBrief model and composition logic
- Business-language title transformation
- Executive summary generation
- Impact area grouping
- Risk extraction and next-steps generation
- Executive Markdown, PDF, and DOCX renderers
- CLI integration for executive audience
"""

from __future__ import annotations

import json
from datetime import date

import pytest
from click.testing import CliRunner

from releasepilot.audience.executive import (
    ExecutiveBrief,
    compose_executive_brief,
)
from releasepilot.cli.app import cli
from releasepilot.domain.enums import Audience, ChangeCategory, Importance
from releasepilot.domain.models import (
    ChangeGroup,
    ChangeItem,
    ReleaseNotes,
    ReleaseRange,
    SourceReference,
)
from releasepilot.rendering.executive_md import ExecutiveMarkdownRenderer


def _pdf_available() -> bool:
    try:
        import reportlab  # noqa: F401

        return True
    except ImportError:
        return False


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


# ── Fixtures ─────────────────────────────────────────────────────────────────


def _make_item(
    title: str,
    category: ChangeCategory = ChangeCategory.FEATURE,
    *,
    breaking: bool = False,
    importance: Importance = Importance.NORMAL,
    description: str = "",
    scope: str = "",
) -> ChangeItem:
    return ChangeItem(
        id=f"id-{title[:10]}",
        title=title,
        description=description,
        category=category,
        scope=scope,
        importance=importance,
        is_breaking=breaking,
        source=SourceReference(commit_hash="abc12345"),
        authors=("alice",),
    )


@pytest.fixture()
def sample_notes() -> ReleaseNotes:
    """ReleaseNotes with a realistic mix of changes."""
    items_by_cat = {
        ChangeCategory.FEATURE: [
            _make_item("feat: add OAuth2 authentication", description="Full OAuth2 flow with PKCE"),
            _make_item("feat: add dark mode support"),
            _make_item("feat: implement search API"),
        ],
        ChangeCategory.BUGFIX: [
            _make_item("fix: pagination off-by-one error", ChangeCategory.BUGFIX),
            _make_item("fix: session token refresh", ChangeCategory.BUGFIX),
        ],
        ChangeCategory.SECURITY: [
            _make_item(
                "fix: patch XSS vulnerability", ChangeCategory.SECURITY, importance=Importance.HIGH
            ),
        ],
        ChangeCategory.PERFORMANCE: [
            _make_item("perf: optimize dashboard queries", ChangeCategory.PERFORMANCE),
        ],
        ChangeCategory.IMPROVEMENT: [
            _make_item("improve error messages", ChangeCategory.IMPROVEMENT),
        ],
        ChangeCategory.BREAKING: [
            _make_item(
                "feat!: remove /api/v1 endpoints",
                ChangeCategory.BREAKING,
                breaking=True,
                description="All v1 API endpoints have been removed",
            ),
        ],
        ChangeCategory.DEPRECATION: [
            _make_item("deprecate legacy webhook format", ChangeCategory.DEPRECATION),
        ],
    }

    groups = tuple(
        ChangeGroup(category=cat, items=tuple(items)) for cat, items in items_by_cat.items()
    )

    all_items = [i for items in items_by_cat.values() for i in items]
    highlights = tuple(i for i in all_items if i.is_breaking or i.importance == Importance.HIGH)
    breaking = tuple(i for i in all_items if i.is_breaking)

    return ReleaseNotes(
        release_range=ReleaseRange(
            from_ref="v2.0.0",
            to_ref="v3.0.0",
            version="3.0.0",
            title="Release 3.0.0",
            release_date=date(2026, 3, 16),
        ),
        groups=groups,
        highlights=highlights,
        breaking_changes=breaking,
        total_changes=len(all_items),
    )


@pytest.fixture()
def sample_brief(sample_notes: ReleaseNotes) -> ExecutiveBrief:
    return compose_executive_brief(sample_notes)


# ── ExecutiveBrief composition tests ─────────────────────────────────────────


class TestComposeExecutiveBrief:
    """Scenarios for composing an executive brief from mixed release notes."""

    def test_returns_executive_brief(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking its type."""

        """THEN it is an ExecutiveBrief instance."""
        assert isinstance(sample_brief, ExecutiveBrief)

    def test_executive_summary_is_non_empty(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking the summary length."""

        """THEN it exceeds 50 characters."""
        assert len(sample_brief.executive_summary) > 50

    def test_summary_mentions_capabilities(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN examining the summary."""

        """THEN it mentions capabilities."""
        assert "capabilit" in sample_brief.executive_summary.lower()

    def test_summary_mentions_security(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN examining the summary."""

        """THEN it mentions security."""
        assert "security" in sample_brief.executive_summary.lower()

    def test_summary_mentions_attention(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN examining the summary."""

        """THEN it mentions attention items."""
        assert "attention" in sample_brief.executive_summary.lower()

    def test_key_achievements_non_empty(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking key achievements."""

        """THEN at least one exists."""
        assert len(sample_brief.key_achievements) >= 1

    def test_key_achievements_max_7(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN counting key achievements."""

        """THEN there are at most seven."""
        assert len(sample_brief.key_achievements) <= 7

    def test_impact_areas_non_empty(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking impact areas."""

        """THEN at least three exist."""
        assert len(sample_brief.impact_areas) >= 3

    def test_impact_areas_have_business_titles(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""
        titles = {a.title for a in sample_brief.impact_areas}

        """WHEN checking impact area titles."""

        """THEN business-oriented titles are used."""
        # Should have business-oriented titles, not technical category names
        assert "New Capabilities" in titles
        assert "Quality & Reliability" in titles

    def test_impact_areas_have_summaries(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking impact area summaries."""

        """THEN each has a meaningful summary."""
        for area in sample_brief.impact_areas:
            assert len(area.summary) > 10

    def test_risks_from_breaking_changes(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking risks."""

        """THEN breaking changes produce API-related risks."""
        assert len(sample_brief.risks) >= 1
        # Should mention API endpoints removal
        any_api_risk = any(
            "api" in r.lower() or "endpoint" in r.lower() for r in sample_brief.risks
        )
        assert any_api_risk

    def test_risks_include_deprecations(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking risks."""

        """THEN deprecations are included."""
        any_deprecation = any("deprecat" in r.lower() for r in sample_brief.risks)
        assert any_deprecation

    def test_next_steps_non_empty(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking next steps."""

        """THEN at least two exist."""
        assert len(sample_brief.next_steps) >= 2

    def test_next_steps_mention_breaking_communication(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking next steps."""

        """THEN breaking change communication is mentioned."""
        any_breaking_step = any("breaking" in s.lower() for s in sample_brief.next_steps)
        assert any_breaking_step

    def test_next_steps_mention_security(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking next steps."""

        """THEN security follow-up is mentioned."""
        any_security_step = any("security" in s.lower() for s in sample_brief.next_steps)
        assert any_security_step

    def test_metrics_has_total_changes(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN reading metrics."""

        """THEN total_changes equals 10."""
        assert sample_brief.metrics["total_changes"] == 10

    def test_metrics_has_features(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN reading metrics."""

        """THEN features count equals 3."""
        assert sample_brief.metrics["features"] == 3

    def test_metrics_has_breaking(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN reading metrics."""

        """THEN breaking count equals 1."""
        assert sample_brief.metrics["breaking"] == 1

    def test_report_title_includes_version(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking the report title."""

        """THEN it includes the version number."""
        assert "3.0.0" in sample_brief.report_title

    def test_report_date_formatted(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking the report date."""

        """THEN it contains the month name."""
        assert "March" in sample_brief.report_date


class TestBusinessLanguageTransformation:
    """Scenarios for business-language transformation of commit titles."""

    def test_conventional_prefix_stripped(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""
        # "feat: add OAuth2 authentication" should become "Add OAuth2 authentication"
        all_items = []
        for area in sample_brief.impact_areas:
            all_items.extend(area.items)

        """WHEN checking impact area items."""

        """THEN conventional commit prefixes are stripped."""
        # No items should start with "feat:", "fix:", etc.
        for item in all_items:
            assert not item.startswith("feat:")
            assert not item.startswith("fix:")
            assert not item.startswith("perf:")

    def test_titles_capitalized(self, sample_brief: ExecutiveBrief):
        """GIVEN a composed executive brief."""

        """WHEN checking item titles."""

        """THEN each title starts with an uppercase letter."""
        for area in sample_brief.impact_areas:
            for item in area.items:
                assert item[0].isupper(), f"Expected capitalized: {item}"


class TestEmptyRelease:
    """Scenarios for composing a brief from empty release notes."""

    def test_empty_notes_produce_valid_brief(self):
        """GIVEN release notes with zero changes."""
        empty_notes = ReleaseNotes(
            release_range=ReleaseRange(from_ref="v1.0", to_ref="HEAD"),
            groups=(),
            total_changes=0,
        )

        """WHEN composing an executive brief."""
        brief = compose_executive_brief(empty_notes)

        """THEN a valid brief with appropriate defaults is produced."""
        assert isinstance(brief, ExecutiveBrief)
        assert "0 changes" in brief.executive_summary or "change" in brief.executive_summary.lower()
        assert len(brief.impact_areas) == 0
        assert len(brief.risks) == 0


# ── Executive Markdown renderer tests ────────────────────────────────────────


class TestExecutiveMarkdownRenderer:
    """Scenarios for rendering an executive brief to Markdown."""

    def test_contains_executive_summary_heading(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN it contains the Executive Summary heading."""
        assert "## Executive Summary" in output

    def test_contains_key_achievements(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN it contains the Key Achievements heading."""
        assert "## Key Achievements" in output

    def test_contains_risks_section(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN it contains the Risks section."""
        assert "## Risks & Attention Items" in output

    def test_contains_next_steps(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN it contains the Next Steps section."""
        assert "## Recommended Next Steps" in output

    def test_contains_metrics_table(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN it contains the metrics table."""
        assert "## Release Metrics" in output
        assert "Total Changes" in output

    def test_contains_report_title(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN it contains the report title with version."""
        assert "Release Brief" in output
        assert "3.0.0" in output

    def test_no_conventional_commit_prefixes(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render(sample_brief)

        """WHEN rendered to Markdown."""

        """THEN no conventional commit prefixes appear."""
        # Should not contain raw commit prefixes
        assert "feat:" not in output
        assert "fix:" not in output
        assert "perf:" not in output


class TestExecutiveJsonRenderer:
    """Scenarios for rendering an executive brief to JSON."""

    def test_valid_json(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        output = ExecutiveMarkdownRenderer().render_json(sample_brief)

        """WHEN rendered to JSON."""
        data = json.loads(output)

        """THEN the output is valid JSON with correct type."""
        assert data["type"] == "executive_brief"

    def test_json_has_all_sections(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""

        """WHEN rendered to JSON."""
        data = json.loads(ExecutiveMarkdownRenderer().render_json(sample_brief))

        """THEN all expected sections are present."""
        assert "executive_summary" in data
        assert "key_achievements" in data
        assert "impact_areas" in data
        assert "risks" in data
        assert "next_steps" in data
        assert "metrics" in data

    def test_json_metrics_match(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""

        """WHEN rendered to JSON."""
        data = json.loads(ExecutiveMarkdownRenderer().render_json(sample_brief))

        """THEN metrics match the original brief."""
        assert data["metrics"]["total_changes"] == 10
        assert data["metrics"]["features"] == 3


# ── Executive PDF renderer tests ─────────────────────────────────────────────


@pytest.mark.skipif(not _pdf_available(), reason="reportlab not installed")
class TestExecutivePdfRenderer:
    """Scenarios for rendering an executive brief to PDF."""

    def test_produces_pdf_bytes(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        from releasepilot.rendering.executive_pdf import ExecutivePdfRenderer

        """WHEN rendered to PDF."""
        data = ExecutivePdfRenderer().render_bytes(sample_brief)

        """THEN valid PDF bytes are produced."""
        assert data[:5] == b"%PDF-"
        assert len(data) > 500

    def test_pdf_has_meaningful_size(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        from releasepilot.rendering.executive_pdf import ExecutivePdfRenderer

        """WHEN rendered to PDF."""
        data = ExecutivePdfRenderer().render_bytes(sample_brief)

        """THEN the output has meaningful size."""
        assert len(data) > 2000


# ── Executive DOCX renderer tests ────────────────────────────────────────────


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
class TestExecutiveDocxRenderer:
    """Scenarios for rendering an executive brief to DOCX."""

    def test_produces_docx_bytes(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        from releasepilot.rendering.executive_docx import ExecutiveDocxRenderer

        """WHEN rendered to DOCX."""
        data = ExecutiveDocxRenderer().render_bytes(sample_brief)

        """THEN valid DOCX bytes are produced."""
        assert data[:2] == b"PK"  # DOCX is a ZIP archive
        assert len(data) > 1000

    def test_docx_reopens_as_document(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        from io import BytesIO

        from docx import Document

        from releasepilot.rendering.executive_docx import ExecutiveDocxRenderer

        """WHEN rendered to DOCX and reopened."""
        data = ExecutiveDocxRenderer().render_bytes(sample_brief)
        doc = Document(BytesIO(data))

        """THEN it contains multiple paragraphs."""
        # Should have multiple paragraphs
        assert len(doc.paragraphs) > 5

    def test_docx_contains_executive_summary(self, sample_brief: ExecutiveBrief):
        """GIVEN an executive brief."""
        from io import BytesIO

        from docx import Document

        from releasepilot.rendering.executive_docx import ExecutiveDocxRenderer

        """WHEN rendered to DOCX."""
        data = ExecutiveDocxRenderer().render_bytes(sample_brief)
        doc = Document(BytesIO(data))

        """THEN the document contains Executive Summary text."""
        # Collect text from paragraphs AND table cells
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        text = "\n".join(parts)
        assert "Executive Summary" in text


# ── CLI integration tests ────────────────────────────────────────────────────


class TestCliExecutiveAudience:
    """Scenarios for CLI integration with executive audience."""

    def test_generate_executive_markdown(self):
        """GIVEN sample changes and executive audience."""
        runner = CliRunner()

        """WHEN generating Markdown via CLI."""
        result = runner.invoke(
            cli,
            [
                "generate",
                "--source-file",
                "examples/sample_changes.json",
                "--audience",
                "executive",
                "--version",
                "3.0.0",
            ],
        )

        """THEN output contains executive sections."""
        assert result.exit_code == 0
        assert "Executive Summary" in result.output
        assert "Key Achievements" in result.output

    def test_generate_executive_json(self):
        """GIVEN sample changes and executive audience."""
        runner = CliRunner()

        """WHEN generating JSON via CLI."""
        result = runner.invoke(
            cli,
            [
                "generate",
                "--source-file",
                "examples/sample_changes.json",
                "--audience",
                "executive",
                "--format",
                "json",
                "--version",
                "3.0.0",
            ],
        )

        """THEN output is valid executive brief JSON."""
        assert result.exit_code == 0
        data = json.loads(result.output)
        assert data["type"] == "executive_brief"

    @pytest.mark.skipif(not _pdf_available(), reason="reportlab not installed")
    def test_export_executive_pdf(self, tmp_path):
        """GIVEN sample changes and executive audience."""
        runner = CliRunner()
        out = str(tmp_path / "brief.pdf")

        """WHEN exporting PDF via CLI."""
        result = runner.invoke(
            cli,
            [
                "export",
                "--source-file",
                "examples/sample_changes.json",
                "--audience",
                "executive",
                "--format",
                "pdf",
                "--version",
                "3.0.0",
                "-o",
                out,
            ],
        )

        """THEN a valid PDF file is created."""
        assert result.exit_code == 0
        assert (tmp_path / "brief.pdf").exists()
        data = (tmp_path / "brief.pdf").read_bytes()
        assert data[:5] == b"%PDF-"

    @pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
    def test_export_executive_docx(self, tmp_path):
        """GIVEN sample changes and executive audience."""
        runner = CliRunner()
        out = str(tmp_path / "brief.docx")

        """WHEN exporting DOCX via CLI."""
        result = runner.invoke(
            cli,
            [
                "export",
                "--source-file",
                "examples/sample_changes.json",
                "--audience",
                "executive",
                "--format",
                "docx",
                "--version",
                "3.0.0",
                "-o",
                out,
            ],
        )

        """THEN a valid DOCX file is created."""
        assert result.exit_code == 0
        assert (tmp_path / "brief.docx").exists()
        data = (tmp_path / "brief.docx").read_bytes()
        assert data[:2] == b"PK"

    def test_executive_no_raw_commit_prefixes(self):
        """GIVEN sample changes and executive audience."""
        runner = CliRunner()

        """WHEN generating via CLI."""
        result = runner.invoke(
            cli,
            [
                "generate",
                "--source-file",
                "examples/sample_changes.json",
                "--audience",
                "executive",
                "--version",
                "3.0.0",
            ],
        )

        """THEN no raw commit prefixes appear."""
        assert result.exit_code == 0
        # Should not contain raw conventional commit prefixes
        assert "feat:" not in result.output
        assert "fix:" not in result.output


# ── Audience view tests ──────────────────────────────────────────────────────


class TestExecutiveAudienceView:
    """Scenarios for executive audience view filtering."""

    def test_executive_view_filters_refactors(self, sample_notes: ReleaseNotes):
        """GIVEN release notes with a refactor group."""
        from releasepilot.audience.views import apply_audience

        # Add a refactor group to sample_notes
        refactor_group = ChangeGroup(
            category=ChangeCategory.REFACTOR,
            items=(_make_item("refactor auth module", ChangeCategory.REFACTOR),),
        )
        notes_with_refactor = ReleaseNotes(
            release_range=sample_notes.release_range,
            groups=sample_notes.groups + (refactor_group,),
            highlights=sample_notes.highlights,
            breaking_changes=sample_notes.breaking_changes,
            total_changes=sample_notes.total_changes + 1,
        )

        """WHEN applying executive audience filter."""
        filtered = apply_audience(notes_with_refactor, Audience.EXECUTIVE)
        categories = {g.category for g in filtered.groups}

        """THEN refactor and infrastructure categories are excluded."""
        assert ChangeCategory.REFACTOR not in categories
        assert ChangeCategory.INFRASTRUCTURE not in categories

    def test_executive_view_keeps_features(self, sample_notes: ReleaseNotes):
        """GIVEN release notes with features."""
        from releasepilot.audience.views import apply_audience

        """WHEN applying executive audience filter."""
        filtered = apply_audience(sample_notes, Audience.EXECUTIVE)
        categories = {g.category for g in filtered.groups}

        """THEN feature category is preserved."""
        assert ChangeCategory.FEATURE in categories

    def test_executive_view_keeps_security(self, sample_notes: ReleaseNotes):
        """GIVEN release notes with security changes."""
        from releasepilot.audience.views import apply_audience

        """WHEN applying executive audience filter."""
        filtered = apply_audience(sample_notes, Audience.EXECUTIVE)
        categories = {g.category for g in filtered.groups}

        """THEN security category is preserved."""
        assert ChangeCategory.SECURITY in categories

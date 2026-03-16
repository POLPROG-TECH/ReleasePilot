"""Narrative DOCX renderer.

Produces a polished, prose-style Word document from a NarrativeBrief.
Unlike the standard DocxRenderer (which produces bullet lists), this renderer
outputs continuous paragraphs suitable for stakeholder communication.

Design mirrors the executive DOCX renderer's visual language but adapts
for narrative prose content rather than structured executive sections.
"""

from __future__ import annotations

import io

from releasepilot.audience.narrative import NarrativeBrief
from releasepilot.rendering import REPO_URL


def _translate(text: str, lang: str) -> str:
    """Translate generated text when the target language is not English."""
    if lang == "en" or not text.strip():
        return text
    try:
        from releasepilot.i18n import translate_text

        return translate_text(text, target_lang=lang)
    except Exception:  # noqa: BLE001
        return text


# ── Design tokens (hex strings for python-docx) ─────────────────────────────

_NAVY = "0F172A"
_SLATE = "334155"
_CHARCOAL = "1E293B"
_TEXT_GRAY = "475569"
_META_GRAY = "64748B"
_LIGHT_GRAY = "94A3B8"
_BORDER = "E2E8F0"
_BG_LIGHT = "F8FAFC"
_WARN_AMBER = "92400E"
_WARN_BG = "FFFBEB"


def _rgb(hex6: str):
    """Create RGBColor from 6-char hex string."""
    from docx.shared import RGBColor

    return RGBColor(int(hex6[:2], 16), int(hex6[2:4], 16), int(hex6[4:], 16))


class NarrativeDocxRenderer:
    """Renders a NarrativeBrief as a polished Word document."""

    def render_bytes(
        self,
        brief: NarrativeBrief,
        *,
        lang: str = "en",
        accent_color: str = "#FB6400",
    ) -> bytes:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        _accent_hex = accent_color.lstrip("#").upper()

        doc = Document()

        # ── Page setup ───────────────────────────────────────────────────
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # ── Default font ────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = _rgb(_CHARCOAL)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = Pt(16)

        rr = brief.release_range
        from releasepilot.i18n import get_label

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # TITLE BLOCK
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if rr.app_name:
            app_para = doc.add_paragraph()
            app_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            app_para.paragraph_format.space_before = Pt(4)
            app_para.paragraph_format.space_after = Pt(4)
            run = app_para.add_run(rr.app_name)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(36)
            run.font.color.rgb = _rgb(_NAVY)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(6)
        tr = title_para.add_run(brief.localized_title(lang))
        tr.font.name = "Calibri"
        tr.font.size = Pt(16)
        tr.font.color.rgb = _rgb(_SLATE)

        # Metadata line
        meta_parts: list[str] = [brief.report_date]
        if rr.version:
            meta_parts.append(f"{get_label('version', lang)} {rr.version}")
        meta_para = doc.add_paragraph()
        meta_para.paragraph_format.space_after = Pt(2)
        mr = meta_para.add_run("  ·  ".join(meta_parts))
        mr.font.name = "Calibri"
        mr.font.size = Pt(9.5)
        mr.font.color.rgb = _rgb(_META_GRAY)

        # Thick accent divider
        _add_accent_rule(doc, qn, _accent_hex, thickness="18")

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # OVERVIEW — highlighted panel
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if brief.mode == "customer-narrative":
            overview_heading_text = get_label("narrative_overview_customer", lang)
        else:
            overview_heading_text = get_label("narrative_overview", lang)

        overview_table = doc.add_table(rows=2, cols=1)
        overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_cell_shading(overview_table.rows[0].cells[0], _BG_LIGHT, qn)
        _set_cell_shading(overview_table.rows[1].cells[0], _BG_LIGHT, qn)

        h_cell = overview_table.rows[0].cells[0]
        h_para = h_cell.paragraphs[0]
        h_para.paragraph_format.space_before = Pt(10)
        h_para.paragraph_format.space_after = Pt(4)
        hr = h_para.add_run(overview_heading_text)
        hr.bold = True
        hr.font.name = "Calibri"
        hr.font.size = Pt(12.5)
        hr.font.color.rgb = _rgb(_SLATE)

        b_cell = overview_table.rows[1].cells[0]
        b_para = b_cell.paragraphs[0]
        b_para.paragraph_format.space_after = Pt(10)
        br_ = b_para.add_run(_translate(brief.overview, lang))
        br_.font.name = "Calibri"
        br_.font.size = Pt(11)
        br_.font.color.rgb = _rgb(_CHARCOAL)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # BODY PARAGRAPHS — continuous prose
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        for paragraph in brief.body_paragraphs:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.line_spacing = Pt(16)
            run = para.add_run(_translate(paragraph, lang))
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.color.rgb = _rgb(_CHARCOAL)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # BREAKING CHANGES — amber warning panel
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if brief.breaking_notice:
            warn_table = doc.add_table(rows=2, cols=1)
            warn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_cell_shading(warn_table.rows[0].cells[0], _WARN_BG, qn)
            _set_cell_shading(warn_table.rows[1].cells[0], _WARN_BG, qn)

            wh_cell = warn_table.rows[0].cells[0]
            wh_para = wh_cell.paragraphs[0]
            wh_para.paragraph_format.space_before = Pt(8)
            wh_para.paragraph_format.space_after = Pt(4)
            whr = wh_para.add_run(get_label("narrative_breaking", lang))
            whr.bold = True
            whr.font.name = "Calibri"
            whr.font.size = Pt(12.5)
            whr.font.color.rgb = _rgb(_WARN_AMBER)

            wb_cell = warn_table.rows[1].cells[0]
            wb_para = wb_cell.paragraphs[0]
            wb_para.paragraph_format.space_after = Pt(8)
            wbr = wb_para.add_run(_translate(brief.breaking_notice, lang))
            wbr.font.name = "Calibri"
            wbr.font.size = Pt(10.5)
            wbr.font.color.rgb = _rgb(_WARN_AMBER)

            _add_table_left_border(warn_table, qn, _WARN_AMBER)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # CLOSING
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if brief.closing:
            _add_thin_rule(doc, qn, _BORDER)
            closing_para = doc.add_paragraph()
            closing_para.paragraph_format.space_after = Pt(8)
            cr = closing_para.add_run(_translate(brief.closing, lang))
            cr.font.italic = True
            cr.font.name = "Calibri"
            cr.font.size = Pt(10)
            cr.font.color.rgb = _rgb(_TEXT_GRAY)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # FOOTER
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        _add_thin_rule(doc, qn, _BORDER)

        from datetime import UTC, datetime

        from releasepilot.i18n import get_label as _gl
        from releasepilot.rendering import AUTHOR, TOOL_NAME
        from releasepilot.rendering.docx_renderer import _add_hyperlink

        now = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_tpl = _gl("footer_generated", lang)
        footer_text = footer_tpl.format(
            tool=TOOL_NAME, author="__AUTHOR__", datetime=now
        )
        parts = footer_text.split("__AUTHOR__")
        fr = footer.add_run(parts[0])
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = _rgb(_LIGHT_GRAY)
        _add_hyperlink(
            footer, REPO_URL, AUTHOR, Pt(7.5), _rgb(_LIGHT_GRAY)
        )
        if len(parts) > 1:
            fr2 = footer.add_run(parts[1])
            fr2.font.size = Pt(7.5)
            fr2.font.color.rgb = _rgb(_LIGHT_GRAY)

        # Provenance note
        provenance = _gl("narrative_provenance", lang).format(
            fact_count=brief.total_facts,
            source_count=len(brief.source_item_ids),
        )
        prov_para = doc.add_paragraph()
        prov_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = prov_para.add_run(provenance)
        pr.font.size = Pt(7.5)
        pr.font.color.rgb = _rgb(_LIGHT_GRAY)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ── Helpers ──────────────────────────────────────────────────────────────────


def _set_cell_shading(cell, hex_color: str, qn) -> None:
    """Apply background shading to a table cell."""
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_accent_rule(doc, qn, color: str, thickness: str = "18") -> None:
    """Add a thick accent-colored horizontal rule."""
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(14)
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): thickness,
            qn("w:space"): "1",
            qn("w:color"): color,
        },
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_thin_rule(doc, qn, color: str) -> None:
    """Add a thin horizontal rule for section separation."""
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): color,
        },
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_table_left_border(table, qn, color: str) -> None:
    """Add a thick left border to a table for accent effect."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    left = borders.makeelement(
        qn("w:left"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "18",
            qn("w:space"): "0",
            qn("w:color"): color,
        },
    )
    borders.append(left)
    tbl_pr.append(borders)

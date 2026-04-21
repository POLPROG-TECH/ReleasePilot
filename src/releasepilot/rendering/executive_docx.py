"""Executive DOCX renderer.

Produces a premium, presentation-ready Word document executive brief designed
for management, leadership, and board audiences.

Design principles mirror the executive PDF renderer:
- Strong title block with generous whitespace
- Accent-bar section headings for scannability
- Highlighted executive summary panel
- Visual metrics dashboard
- Clear visual hierarchy and breathing room throughout
- Risk section with distinct visual treatment
- Minimal, elegant footer
"""

from __future__ import annotations

import io

from releasepilot.audience.executive import ExecutiveBrief
from releasepilot.rendering import REPO_URL


def _translate(text: str, lang: str) -> str:
    """Translate generated text when the target language is not English."""
    if lang == "en" or not text.strip():
        return text
    try:
        from releasepilot.i18n import translate_text

        return translate_text(text, target_lang=lang)
    except Exception:  # noqa: BLE001
        return text


# ── Design tokens (hex strings for XML, RGBColor for python-docx) ─────────

_NAVY = "0F172A"
_SLATE = "334155"
_CHARCOAL = "1E293B"
_TEXT_GRAY = "475569"
_META_GRAY = "64748B"
_LIGHT_GRAY = "94A3B8"
_BORDER = "E2E8F0"
_BG_LIGHT = "F8FAFC"
_BG_METRICS = "F1F5F9"
_ACCENT = "3B82F6"
_RISK_RED = "B91C1C"
_RISK_BG = "FEF2F2"
_RISK_BORDER = "FCA5A5"


def _rgb(hex6: str):
    """Create RGBColor from 6-char hex string."""
    from docx.shared import RGBColor

    return RGBColor(int(hex6[:2], 16), int(hex6[2:4], 16), int(hex6[4:], 16))


class ExecutiveDocxRenderer:
    """Renders an ExecutiveBrief as a premium Word document."""

    def render_bytes(
        self, brief: ExecutiveBrief, *, lang: str = "en", accent_color: str = "#FB6400"
    ) -> bytes:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        # Derive 6-char hex from accent_color parameter
        _accent_hex = accent_color.lstrip("#").upper()

        doc = Document()

        # ── Page setup ───────────────────────────────────────────────────
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # ── Default font ────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = _rgb(_CHARCOAL)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = Pt(16)

        rr = brief.release_range
        from releasepilot.i18n import get_label

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # TITLE BLOCK
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if rr.app_name:
            app_para = doc.add_paragraph()
            app_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            app_para.paragraph_format.space_before = Pt(4)
            app_para.paragraph_format.space_after = Pt(4)
            run = app_para.add_run(rr.app_name)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(36)
            run.font.color.rgb = _rgb(_NAVY)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(6)
        tr = title_para.add_run(brief.localized_title(lang))
        tr.font.name = "Calibri"
        tr.font.size = Pt(16)
        tr.font.color.rgb = _rgb(_SLATE)

        # Metadata line
        meta_parts: list[str] = [brief.localized_date(lang)]
        if rr.version:
            meta_parts.append(f"{get_label('version', lang)} {rr.version}")
        if brief.analysis_period:
            meta_parts.append(
                get_label("analysis_period", lang).format(
                    period=brief.analysis_period,
                ),
            )
        meta_para = doc.add_paragraph()
        meta_para.paragraph_format.space_after = Pt(2)
        mr = meta_para.add_run("  ·  ".join(meta_parts))
        mr.font.name = "Calibri"
        mr.font.size = Pt(9.5)
        mr.font.color.rgb = _rgb(_META_GRAY)

        # Thick accent divider
        _add_accent_rule(doc, qn, _accent_hex, thickness="18")

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # EXECUTIVE SUMMARY - highlighted panel
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        summary_table = doc.add_table(rows=2, cols=1)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_cell_shading(summary_table.rows[0].cells[0], _BG_LIGHT, qn)
        _set_cell_shading(summary_table.rows[1].cells[0], _BG_LIGHT, qn)

        # Heading inside panel
        h_cell = summary_table.rows[0].cells[0]
        h_para = h_cell.paragraphs[0]
        h_para.paragraph_format.space_before = Pt(10)
        h_para.paragraph_format.space_after = Pt(4)
        hr = h_para.add_run(get_label("executive_summary", lang))
        hr.bold = True
        hr.font.name = "Calibri"
        hr.font.size = Pt(12.5)
        hr.font.color.rgb = _rgb(_SLATE)

        # Body inside panel
        b_cell = summary_table.rows[1].cells[0]
        b_para = b_cell.paragraphs[0]
        b_para.paragraph_format.space_after = Pt(10)
        br_ = b_para.add_run(_translate(brief.executive_summary, lang))
        br_.font.name = "Calibri"
        br_.font.size = Pt(11)
        br_.font.color.rgb = _rgb(_CHARCOAL)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # METRICS DASHBOARD
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        metrics_rows = _build_metrics_rows(brief, lang)
        if len(metrics_rows) > 1:
            _add_accent_section_heading(
                doc,
                qn,
                get_label("release_metrics", lang),
                _accent_hex,
                _SLATE,
            )

            table = doc.add_table(rows=len(metrics_rows), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Cm(10)
            table.columns[1].width = Cm(3)

            # Header row
            for j, val in enumerate(metrics_rows[0]):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(9.5)
                r.font.color.rgb = _rgb(_SLATE)
                _set_cell_shading(cell, _BG_METRICS, qn)
                if j == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Data rows
            for i, row_data in enumerate(metrics_rows[1:], 1):
                for j, val in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    r = p.add_run(val)
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                    r.font.color.rgb = _rgb(_CHARCOAL)
                    if j == 1:
                        r.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Bottom border on last row
            _add_table_bottom_border(table, qn, _BORDER)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # KEY ACHIEVEMENTS
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if brief.key_achievements:
            _add_accent_section_heading(
                doc,
                qn,
                get_label("key_achievements", lang),
                _accent_hex,
                _SLATE,
            )
            for i, item in enumerate(brief.key_achievements, 1):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(16)
                para.paragraph_format.space_after = Pt(5)
                nr = para.add_run(f"{i}.  ")
                nr.bold = True
                nr.font.color.rgb = _rgb(_SLATE)
                ir = para.add_run(item)
                ir.font.color.rgb = _rgb(_CHARCOAL)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # IMPACT AREAS
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        for area in brief.impact_areas:
            _add_accent_section_heading(
                doc,
                qn,
                _translate(area.title, lang),
                _accent_hex,
                _SLATE,
            )
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(6)
            sp.paragraph_format.left_indent = Pt(4)
            sr = sp.add_run(_translate(area.summary, lang))
            sr.font.italic = True
            sr.font.size = Pt(10)
            sr.font.color.rgb = _rgb(_TEXT_GRAY)
            for item in area.items:
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Pt(16)
                bp.paragraph_format.space_after = Pt(4)
                bp.add_run(f"•  {item}").font.color.rgb = _rgb(_CHARCOAL)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # RISKS - distinct visual treatment
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if brief.risks:
            risk_table = doc.add_table(rows=1 + len(brief.risks), cols=1)
            risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Heading row
            hc = risk_table.rows[0].cells[0]
            _set_cell_shading(hc, _RISK_BG, qn)
            hp = hc.paragraphs[0]
            hp.paragraph_format.space_before = Pt(8)
            hp.paragraph_format.space_after = Pt(4)
            rhr = hp.add_run(get_label("risks_attention", lang))
            rhr.bold = True
            rhr.font.name = "Calibri"
            rhr.font.size = Pt(12.5)
            rhr.font.color.rgb = _rgb(_RISK_RED)

            for i, risk in enumerate(brief.risks):
                rc = risk_table.rows[i + 1].cells[0]
                _set_cell_shading(rc, _RISK_BG, qn)
                rp = rc.paragraphs[0]
                rp.paragraph_format.space_after = Pt(4)
                rp.paragraph_format.left_indent = Pt(8)
                rr_ = rp.add_run(f"•  {risk}")
                rr_.font.color.rgb = _rgb(_RISK_RED)
                rr_.font.size = Pt(10.5)

            # Left accent border on the table
            _add_table_left_border(risk_table, qn, _RISK_RED)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # NEXT STEPS
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        if brief.next_steps:
            _add_accent_section_heading(
                doc,
                qn,
                get_label("next_steps", lang),
                _accent_hex,
                _SLATE,
            )
            for step in brief.next_steps:
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Pt(16)
                bp.paragraph_format.space_after = Pt(4)
                bp.add_run(f"•  {_translate(step, lang)}").font.color.rgb = _rgb(_CHARCOAL)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        # FOOTER
        # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
        _add_thin_rule(doc, qn, _BORDER)

        from datetime import UTC, datetime

        from releasepilot.i18n import get_label as _gl
        from releasepilot.rendering import AUTHOR, TOOL_NAME
        from releasepilot.rendering.docx_renderer import _add_hyperlink

        now = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_tpl = _gl("footer_generated", lang)
        footer_text = footer_tpl.format(
            tool=TOOL_NAME,
            author="__AUTHOR__",
            datetime=now,
        )
        parts = footer_text.split("__AUTHOR__")
        fr = footer.add_run(parts[0])
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = _rgb(_LIGHT_GRAY)
        _add_hyperlink(
            footer,
            REPO_URL,
            AUTHOR,
            Pt(7.5),
            _rgb(_LIGHT_GRAY),
        )
        if len(parts) > 1:
            fr2 = footer.add_run(parts[1])
            fr2.font.size = Pt(7.5)
            fr2.font.color.rgb = _rgb(_LIGHT_GRAY)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ── Helpers ──────────────────────────────────────────────────────────────────


def _build_metrics_rows(
    brief: ExecutiveBrief,
    lang: str = "en",
) -> list[list[str]]:
    """Build translated metrics rows for the dashboard table."""
    from releasepilot.i18n import get_label

    m = brief.metrics
    rows = [[get_label("metric", lang), get_label("value", lang)]]
    rows.append([get_label("total_changes", lang), str(m.get("total_changes", 0))])
    for key, lk in (
        ("features", "new_features"),
        ("improvements", "improvements"),
        ("bugfixes", "issues_resolved"),
        ("performance", "performance_gains"),
        ("security", "security_fixes"),
    ):
        val = m.get(key, 0)
        if val:
            rows.append([get_label(lk, lang), str(val)])
    if m.get("breaking", 0):
        rows.append([get_label("breaking_changes", lang), str(m["breaking"])])
    return rows


def _set_cell_shading(cell, hex_color: str, qn) -> None:
    """Apply background shading to a table cell."""
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_accent_rule(doc, qn, color: str, thickness: str = "18") -> None:
    """Add a thick accent-colored horizontal rule."""
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(14)
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): thickness,
            qn("w:space"): "1",
            qn("w:color"): color,
        },
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_thin_rule(doc, qn, color: str) -> None:
    """Add a thin horizontal rule for footer separation."""
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): color,
        },
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_accent_section_heading(
    doc,
    qn,
    text: str,
    accent_color: str,
    text_color: str,
) -> None:
    """Add a section heading with a left accent bar for scannability."""
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12.5)
    run.font.color.rgb = _rgb(text_color)

    # Left border via paragraph border XML
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    left = p_bdr.makeelement(
        qn("w:left"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "24",
            qn("w:space"): "8",
            qn("w:color"): accent_color,
        },
    )
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _add_table_bottom_border(table, qn, color: str) -> None:
    """Add a bottom border to the last row of a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    bottom = borders.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): color,
        },
    )
    borders.append(bottom)
    tbl_pr.append(borders)


def _add_table_left_border(table, qn, color: str) -> None:
    """Add a thick left border to a table for accent effect."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    left = borders.makeelement(
        qn("w:left"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "18",
            qn("w:space"): "0",
            qn("w:color"): color,
        },
    )
    borders.append(left)
    tbl_pr.append(borders)

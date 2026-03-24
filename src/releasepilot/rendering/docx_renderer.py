"""DOCX renderer for ReleasePilot.

Produces polished, professional Word documents suitable for stakeholder
communication, release handoffs, and archival documentation.

Uses python-docx for document generation.
"""

from __future__ import annotations

import io

from releasepilot.config.settings import RenderConfig
from releasepilot.domain.models import ChangeItem, ReleaseNotes
from releasepilot.rendering import REPO_URL

# Emoji → text label mapping for clean DOCX output
_EMOJI_STRIP = {
    "⚠️": "",
    "🔒": "",
    "✨": "",
    "🔧": "",
    "🐛": "",
    "⚡": "",
    "📦": "",
    "📝": "",
    "🏗️": "",
    "♻️": "",
    "📋": "",
}


def _clean_label(text: str) -> str:
    """Remove emojis from display labels for DOCX."""
    for emoji in _EMOJI_STRIP:
        text = text.replace(emoji, "")
    return text.strip()


def _translate_label(text: str, lang: str) -> str:
    """Translate a display label when the target language is not English."""
    if lang == "en" or not text.strip():
        return text
    try:
        from releasepilot.i18n import translate_text

        return translate_text(text, target_lang=lang)
    except Exception:  # noqa: BLE001
        return text


class DocxRenderer:
    """Renders ReleaseNotes as a professional Word document."""

    def render(self, notes: ReleaseNotes, config: RenderConfig) -> str:
        """DOCX is binary — use render_bytes() instead.

        Raises NotImplementedError to prevent silent empty-string returns.
        """
        raise NotImplementedError(
            "DocxRenderer.render() is not supported. Use render_bytes() for DOCX output."
        )

    def render_bytes(self, notes: ReleaseNotes, config: RenderConfig) -> bytes:
        """Render release notes to a DOCX byte buffer."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        # Accent color from config (strip leading '#')
        accent_hex = config.accent_color.lstrip("#")

        doc = Document()

        # ── Page setup ──
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # ── Default font ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        rr = notes.release_range
        lang = config.language

        # ── Title block — app name centered on its own, subtitle below ──
        if rr.app_name:
            app_para = doc.add_heading(rr.app_name, level=0)
            app_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in app_para.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                run.font.size = Pt(28)
            title_para = doc.add_heading(rr.subtitle, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in title_para.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                run.font.size = Pt(20)
        else:
            title_para = doc.add_heading(rr.display_title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in title_para.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                run.font.size = Pt(24)

        # ── Subtitle / metadata ──
        subtitle_parts = []
        if rr.version:
            from releasepilot.i18n import get_label as _label

            subtitle_parts.append(f"{_label('version', lang)} {rr.version}")
        if rr.release_date:
            from releasepilot.i18n import get_label as _label

            released = _label("released_on", lang).format(date=rr.release_date.isoformat())
            subtitle_parts.append(released)
        from releasepilot.i18n import get_label as _label

        subtitle_parts.append(_label("changes_in_release", lang).format(count=notes.total_changes))
        if notes.metadata.get("pipeline_summary"):
            subtitle_parts.append(notes.metadata["pipeline_summary"])

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta.space_after = Pt(12)
        run = meta.add_run(" · ".join(subtitle_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.italic = True

        # ── Horizontal rule ──
        _add_horizontal_rule(doc, accent_hex)

        # ── Highlights ──
        if notes.highlights:
            from releasepilot.i18n import get_label as _label

            h = doc.add_heading(_label("highlights", lang), level=1)
            _style_heading(h)
            for item in notes.highlights:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(item.title)
                run.bold = True
                if item.description:
                    para.add_run(f" — {item.description[:200]}")

        # ── Breaking Changes ──
        if notes.breaking_changes:
            from releasepilot.i18n import get_label as _label

            h = doc.add_heading(_label("breaking_changes", lang), level=1)
            _style_heading(h, color=RGBColor(0xC0, 0x39, 0x2B))
            for item in notes.breaking_changes:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(item.title)
                run.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                if item.description:
                    desc_run = para.add_run(f"\n{item.description[:300]}")
                    desc_run.font.size = Pt(9)
                    desc_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # ── Category groups ──
        breaking_ids = {item.id for item in notes.breaking_changes}
        for group in notes.groups:
            if group.category.value == "breaking" and notes.breaking_changes:
                continue

            label = _clean_label(group.display_label)
            label = _translate_label(label, lang)
            h = doc.add_heading(label, level=1)
            _style_heading(h)

            for item in group.items:
                if item.id in breaking_ids:
                    continue

                para = doc.add_paragraph(style="List Bullet")
                para.add_run(item.title)
                suffix = _item_suffix_text(item, config)
                if suffix:
                    run = para.add_run(suffix)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ── Footer ──
        from datetime import UTC, datetime

        from releasepilot.i18n import get_label as _label
        from releasepilot.rendering import AUTHOR, TOOL_NAME

        now = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_tpl = _label("footer_generated", lang)
        # Split the template around the author placeholder to insert a hyperlink
        footer_text_before = footer_tpl.format(tool=TOOL_NAME, author="__AUTHOR__", datetime=now)
        parts = footer_text_before.split("__AUTHOR__")
        run = footer.add_run(parts[0])
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        _add_hyperlink(footer, REPO_URL, AUTHOR, Pt(7), RGBColor(0x88, 0xAA, 0xCC))
        if len(parts) > 1:
            run2 = footer.add_run(parts[1])
            run2.font.size = Pt(7)
            run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


def _style_heading(heading, *, color=None) -> None:
    """Apply consistent styling to section headings."""
    from docx.shared import Pt, RGBColor

    default_color = color or RGBColor(0x16, 0x21, 0x3E)
    for run in heading.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = default_color


def _add_horizontal_rule(doc, color: str = "FB6400") -> None:
    """Add a thin horizontal rule between sections."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    para = doc.add_paragraph()
    para.space_before = Pt(4)
    para.space_after = Pt(8)
    # Create a bottom border on the paragraph
    p_pr = para._p.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): color,
        },
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_hyperlink(paragraph, url: str, text: str, font_size, font_color) -> None:
    """Add a clickable hyperlink run to a python-docx paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Create the w:hyperlink element
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a run inside the hyperlink
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # Font size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size.pt * 2)))
    r_pr.append(sz)
    # Font color
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{font_color[0]:02X}{font_color[1]:02X}{font_color[2]:02X}")
    r_pr.append(color_el)
    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)

    # Text
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _item_suffix_text(item: ChangeItem, config: RenderConfig) -> str:
    """Build inline metadata suffix for an item (plain text)."""
    parts = []
    if config.show_scope and item.scope:
        parts.append(f"[{item.scope}]")
    if config.show_authors and item.authors:
        parts.append(f"by {', '.join(item.authors)}")
    if config.show_commit_hashes and item.source.short_hash:
        parts.append(item.source.short_hash)
    if parts:
        return " — " + " ".join(parts)
    return ""
